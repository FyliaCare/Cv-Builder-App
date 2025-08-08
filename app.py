"""
app.py — Streamlit AI-style CV Builder (full-featured, .docx export)

Dependencies:
  pip install streamlit python-docx pillow

Run:
  streamlit run app.py
"""

import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import random
import base64
from html import escape

st.set_page_config(page_title="✨ CV Studio", layout="wide", initial_sidebar_state="expanded")

# --------------------------
# Helper functions & local "AI"
# --------------------------
ACTION_VERBS = [
    "Led", "Designed", "Implemented", "Spearheaded", "Managed",
    "Optimized", "Increased", "Reduced", "Improved", "Delivered",
    "Built", "Coordinated", "Developed", "Negotiated", "Streamlined",
    "Drove", "Facilitated", "Executed", "Mentored", "Launched"
]
METRICS = ["revenue", "efficiency", "customer satisfaction", "cost", "uptime", "retention", "conversion rate"]


def smart_expand(description: str, role: str = "", company: str = "", n: int = 3):
    """
    Local 'AI' that expands a short description into n achievement-focused bullets.
    Deterministic-ish but randomized for variety.
    """
    if not (description and description.strip()):
        return []

    desc = " ".join(description.strip().split())
    bullets = []

    # metric-based bullet
    metric = random.choice(METRICS)
    value = random.choice([8, 10, 12, 15, 20, 25, 30])
    b1 = f"{random.choice(ACTION_VERBS)} {desc} at {company}." if company else f"{random.choice(ACTION_VERBS)} {desc}."
    b2 = f"{random.choice(ACTION_VERBS)} {desc}, achieving ~{value}% improvement in {metric}."
    b3 = f"{random.choice(ACTION_VERBS)} {desc} by focusing on stakeholder needs and measurable KPIs."

    candidates = [b2, b1, b3]
    # add role-specific flavor
    role_l = (role or "").lower()
    if any(x in role_l for x in ["sales", "account", "business"]):
        candidates.append("Built strong client relationships and expanded accounts through consultative selling.")
    if any(x in role_l for x in ["engineer", "developer", "dev", "software"]):
        candidates.append("Improved system reliability and deployment velocity through automation and testing.")
    if any(x in role_l for x in ["product", "pm", "product manager"]):
        candidates.append("Prioritised features and worked cross-functionally to launch product improvements.")
    random.shuffle(candidates)

    for c in candidates:
        if len(bullets) >= n:
            break
        # ensure snappy phrasing
        s = c.strip()
        if not s.endswith("."):
            s += "."
        bullets.append(s)
    # fallback duplication avoided
    if len(bullets) < n:
        while len(bullets) < n:
            bullets.append(f"{random.choice(ACTION_VERBS)} {desc}.")
    return bullets[:n]


def pil_to_datauri(uploaded_file):
    """Convert an uploaded file (UploadedFile) to a data URI for embedding in HTML."""
    if not uploaded_file:
        return None
    try:
        img = Image.open(uploaded_file).convert("RGB")
        buf = BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        return f"data:image/png;base64,{b64}"
    except Exception:
        return None


# --------------------------
# Session state defaults
# --------------------------
if "profile" not in st.session_state:
    st.session_state.profile = {
        "name": "Jane Doe",
        "title": "Senior Technical Specialist",
        "email": "",
        "phone": "",
        "location": "",
        "linkedin": "",
        "portfolio": "",
        "summary": "Results-driven professional with strong experience in operations and client-facing technical services."
    }
if "experience_list" not in st.session_state:
    # sample initial experience
    st.session_state.experience_list = [
        {
            "role": "Sales Representative",
            "company": "Intertek Geronimo Oil and Gas",
            "period": "2022 — Present",
            "description": "Managed key accounts and supported inspection services",
            "bullets": smart_expand("Managed key accounts and supported inspection services", role="Sales Representative", company="Intertek", n=4)
        }
    ]
if "education_list" not in st.session_state:
    st.session_state.education_list = [
        {"degree": "B.Sc. Mechanical Engineering", "school": "University of Example", "year": "2018"}
    ]
if "skills" not in st.session_state:
    st.session_state.skills = ["Client Relationships", "Inspection", "Asset Integrity"]
if "photo_file" not in st.session_state:
    st.session_state.photo_file = None
if "design" not in st.session_state:
    st.session_state.design = {"style": "Modern Color", "accent": "#0b6efd", "include_photo": True}

# --------------------------
# Page layout (header)
# --------------------------
st.markdown(
    """
    <style>
      .big-title { font-size:32px; font-weight:700; color:#102A43; }
      .muted { color:#516B7A; }
      .app-header { display:flex; align-items:center; gap:18px; }
      .spark { font-size:28px; }
    </style>
    """,
    unsafe_allow_html=True,
)
header_col1, header_col2 = st.columns([4, 1])
with header_col1:
    st.markdown("<div class='app-header'><div class='spark'>✨</div><div><div class='big-title'>CV Studio</div><div class='muted'>AI-style bullets — DOCX export — Beautiful templates</div></div></div>", unsafe_allow_html=True)
with header_col2:
    if st.button("🎯 Fill sample data"):
        # Fill with helpful sample data to demo quickly (includes contact + photo cleared)
        st.session_state.profile = {
            "name": "Jojo Montford",
            "title": "Senior Sales Representative — Inspection Services",
            "email": "jojo@example.com",
            "phone": "+233 123 456 789",
            "location": "Accra, Ghana",
            "linkedin": "https://linkedin.com/in/jojo",
            "portfolio": "",
            "summary": "Sales-driven professional with experience in selling inspection and asset integrity services. Skilled at building client relationships and improving account revenue."
        }
        st.session_state.experience_list = [
            {
                "role": "Sales Representative",
                "company": "Intertek Geronimo Oil and Gas",
                "period": "2021 — Present",
                "description": "Managed regional accounts and closed inspection contracts",
                "bullets": smart_expand("Managed regional accounts and closed inspection contracts", "Sales Representative", "Intertek", 4)
            }
        ]
        st.session_state.education_list = [{"degree": "HND Mechanical Engineering", "school": "Accra Technical University", "year": "2016"}]
        st.session_state.skills = ["Sales", "Client Management", "NDT", "Asset Integrity"]
        st.session_state.photo_file = None
        st.success("Sample data loaded — explore tabs to edit!")

# --------------------------
# Left column: Tabs for input
# --------------------------
left, right = st.columns([2.2, 1])
with left:
    tabs = st.tabs(["👤 Profile", "💼 Experience", "🎓 Education", "🛠 Skills", "🎨 Design", "✅ Preview & Export"])

    # ---------- PROFILE ----------
    with tabs[0]:
        st.subheader("Profile Information")

        col1, col2 = st.columns([1, 2])

        # Photo uploader
        with col1:
            profile_photo = st.file_uploader("Upload Profile Photo", type=["png", "jpg", "jpeg"])
            if profile_photo:
                st.session_state.photo_file = profile_photo

        # Profile details
        with col2:
            full_name = st.text_input("Full Name", value=st.session_state.profile.get("name", ""))
            title = st.text_input("Professional Title", value=st.session_state.profile.get("title", ""), placeholder="e.g. Mechanical Engineer")
            email = st.text_input("Email", value=st.session_state.profile.get("email", ""))
            phone = st.text_input("Phone Number", value=st.session_state.profile.get("phone", ""))
            location = st.text_input("Location", value=st.session_state.profile.get("location", ""), placeholder="City, Country")
            linkedin = st.text_input("LinkedIn URL", value=st.session_state.profile.get("linkedin", ""), placeholder="https://linkedin.com/in/...")
            portfolio = st.text_input("Portfolio/Website", value=st.session_state.profile.get("portfolio", ""), placeholder="https://...")
            summary = st.text_area("Professional Summary", value=st.session_state.profile.get("summary", ""), height=100,
                                   placeholder="Write 3–4 sentences summarizing your career highlights...")

            # Save into session state so preview & export pick it up
            st.session_state.profile.update({
                "name": full_name,
                "title": title,
                "email": email,
                "phone": phone,
                "location": location,
                "linkedin": linkedin,
                "portfolio": portfolio,
                "summary": summary
            })

    # ---------- EXPERIENCE ----------
    with tabs[1]:
        st.header("Experience — add & edit")
        with st.expander("Add a new role", expanded=True):
            exp_role = st.text_input("Role (e.g. Sales Representative)", key="new_role")
            exp_company = st.text_input("Company", key="new_company")
            exp_period = st.text_input("Period (e.g. 2020 — Present)", key="new_period")
            exp_desc = st.text_area("Short description or responsibilities (one or more lines)", key="new_desc", height=90)
            c1, c2, c3 = st.columns([1, 1, 1])
            if c1.button("✨ Add (AI bullets)"):
                bullets = smart_expand(exp_desc or exp_role, role=exp_role, company=exp_company, n=4)
                st.session_state.experience_list.insert(0, {
                    "role": exp_role or "Role",
                    "company": exp_company or "",
                    "period": exp_period or "",
                    "description": exp_desc or "",
                    "bullets": bullets
                })
                # clear inputs
                st.session_state.new_role = ""
                st.session_state.new_company = ""
                st.session_state.new_period = ""
                st.session_state.new_desc = ""
                st.success("Entry added with AI-style bullets")
            if c2.button("➕ Add raw"):
                raw_lines = [ln.strip() for ln in (exp_desc or "").split("\n") if ln.strip()]
                bullets = [ln if ln.endswith(".") else ln + "." for ln in raw_lines] if raw_lines else []
                st.session_state.experience_list.insert(0, {
                    "role": exp_role or "Role",
                    "company": exp_company or "",
                    "period": exp_period or "",
                    "description": exp_desc or "",
                    "bullets": bullets
                })
                st.session_state.new_role = ""
                st.session_state.new_company = ""
                st.session_state.new_period = ""
                st.session_state.new_desc = ""
                st.experimental_rerun()
            if c3.button("Clear all experience"):
                st.session_state.experience_list = []
                st.success("All experiences removed")

        st.markdown("**Existing experience (top = newest). Click edit to change bullets or remove.**")
        # list entries with edit/remove
        for i, ex in enumerate(list(st.session_state.experience_list)):
            with st.expander(f"{ex.get('role','Role')} — {ex.get('company','')}", expanded=False):
                r = st.text_input("Role", value=ex.get("role", ""), key=f"role_{i}")
                c = st.text_input("Company", value=ex.get("company", ""), key=f"company_{i}")
                p = st.text_input("Period", value=ex.get("period", ""), key=f"period_{i}")
                desc_field = st.text_area("Description (notes)", value=ex.get("description", ""), key=f"desc_{i}", height=70)

                # editable bullets
                st.markdown("**Bullets — edit freely**")
                new_bullets = []
                bullets_current = ex.get("bullets", [])
                for bi, b in enumerate(bullets_current):
                    nb = st.text_area(f"Bullet #{bi+1}", value=b, key=f"bullet_{i}_{bi}", height=60)
                    new_bullets.append(nb)
                if st.button("Regenerate bullets (AI) for this role", key=f"regen_{i}"):
                    new_bullets = smart_expand(desc_field or r, role=r, company=c, n=4)
                    st.session_state.experience_list[i]["bullets"] = new_bullets
                    st.experimental_rerun()
                cols = st.columns([1, 1, 1])
                if cols[0].button("Save changes", key=f"save_exp_{i}"):
                    st.session_state.experience_list[i]["role"] = r
                    st.session_state.experience_list[i]["company"] = c
                    st.session_state.experience_list[i]["period"] = p
                    st.session_state.experience_list[i]["description"] = desc_field
                    st.session_state.experience_list[i]["bullets"] = new_bullets
                    st.success("Saved")
                if cols[1].button("Remove entry", key=f"remove_exp_{i}"):
                    st.session_state.experience_list.pop(i)
                    st.experimental_rerun()

    # ---------- EDUCATION ----------
    with tabs[2]:
        st.header("Education")
        with st.expander("Add education", expanded=True):
            ed_degree = st.text_input("Degree / Qualification", key="ed_deg")
            ed_school = st.text_input("School / Institution", key="ed_sch")
            ed_year = st.text_input("Year(s)", key="ed_year")
            if st.button("➕ Add education"):
                st.session_state.education_list.insert(0, {"degree": ed_degree or "Degree", "school": ed_school or "", "year": ed_year or ""})
                st.session_state.ed_deg = ""
                st.session_state.ed_sch = ""
                st.session_state.ed_year = ""
                st.success("Education added")
        if st.session_state.education_list:
            st.markdown("**Existing education**")
            for i, ed in enumerate(list(st.session_state.education_list)):
                cols = st.columns([4, 1])
                with cols[0]:
                    st.write(f"**{ed.get('degree','')}** — {ed.get('school','')} ({ed.get('year','')})")
                with cols[1]:
                    if st.button("Remove", key=f"rem_ed_{i}"):
                        st.session_state.education_list.pop(i)
                        st.experimental_rerun()

    # ---------- SKILLS ----------
    with tabs[3]:
        st.header("Skills")
        sk = st.text_input("Add a skill", key="skill_input")
        if st.button("➕ Add skill"):
            v = sk.strip()
            if v:
                st.session_state.skills.insert(0, v)
                st.session_state.skill_input = ""
                st.success("Skill added")
        if st.session_state.skills:
            st.markdown("**Current skills**")
            for i, s in enumerate(list(st.session_state.skills)):
                cols = st.columns([4, 1])
                with cols[0]:
                    st.write(s)
                with cols[1]:
                    if st.button("Remove", key=f"rem_skill_{i}"):
                        st.session_state.skills.pop(i)
                        st.experimental_rerun()

    # ---------- DESIGN / SETTINGS ----------
    with tabs[4]:
        st.header("Design & Settings")
        style = st.selectbox("Choose style for export/preview", ["Modern Color", "Classic B/W", "Minimal One-Page"])
        accent = st.color_picker("Accent color (for Modern)", st.session_state.design.get("accent", "#0b6efd"))
        include_photo = st.checkbox("Include profile photo in export", value=st.session_state.design.get("include_photo", True))
        st.session_state.design = {"style": style, "accent": accent, "include_photo": include_photo}

    # ---------- PREVIEW & EXPORT ----------
    with tabs[5]:
        st.header("Preview & Download")
        st.markdown("Use the preview on the right to check layout. When ready, download as **.docx** (editable) or HTML.")

        download_col1, download_col2 = st.columns([1, 1])
        if download_col1.button("📥 Build & Show DOCX (prepare for download)"):
            # Build .docx from current session state and present download button
            doc_bytes = BytesIO()
            doc = Document()

            # Header (name + title)
            name = st.session_state.profile.get("name", "")
            title = st.session_state.profile.get("title", "")
            doc.add_heading(name or "Unnamed", level=0)
            if title:
                doc.add_paragraph(title)

            # Contact meta (single line)
            contact_parts = []
            if st.session_state.profile.get("email"):
                contact_parts.append(st.session_state.profile["email"])
            if st.session_state.profile.get("phone"):
                contact_parts.append(st.session_state.profile["phone"])
            if st.session_state.profile.get("location"):
                contact_parts.append(st.session_state.profile["location"])
            if st.session_state.profile.get("linkedin"):
                contact_parts.append(st.session_state.profile["linkedin"])
            if st.session_state.profile.get("portfolio"):
                contact_parts.append(st.session_state.profile["portfolio"])
            if contact_parts:
                doc.add_paragraph(" | ".join(contact_parts))

            # Summary
            if st.session_state.profile.get("summary"):
                doc.add_paragraph(st.session_state.profile.get("summary"))

            # Add photo if included
            if st.session_state.design.get("include_photo", True) and st.session_state.photo_file:
                try:
                    # doc.add_picture expects a file-like object positioned at 0
                    img_bytes = BytesIO(st.session_state.photo_file.getvalue())
                    img_bytes.seek(0)
                    doc.add_picture(img_bytes, width=Inches(1.2))
                except Exception:
                    # silently ignore photo errors
                    pass

            # Education
            if st.session_state.education_list:
                doc.add_heading("Education", level=1)
                for ed in st.session_state.education_list:
                    p = doc.add_paragraph()
                    p.add_run(ed.get("degree", "") + " — ").bold = True
                    p.add_run(ed.get("school", ""))
                    if ed.get("year"):
                        p.add_run(f" ({ed['year']})")

            # Experience
            if st.session_state.experience_list:
                doc.add_heading("Experience", level=1)
                for ex in st.session_state.experience_list:
                    p = doc.add_paragraph()
                    p.add_run(f"{ex.get('role','')} — {ex.get('company','')}").bold = True
                    if ex.get("period"):
                        p.add_run(f"  ({ex.get('period')})")
                    # bullets
                    for b in ex.get("bullets", []):
                        try:
                            doc.add_paragraph(b, style="List Bullet")
                        except Exception:
                            # fallback if style name not present
                            doc.add_paragraph(f"• {b}")

            # Skills
            if st.session_state.skills:
                doc.add_heading("Skills", level=1)
                doc.add_paragraph(", ".join(st.session_state.skills))

            # Save
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            st.session_state._last_docx = doc_bytes.getvalue()
            st.success("DOCX built — use the download button below to save it locally.")

        # provide download button if doc ready
        if "_last_docx" in st.session_state:
            st.download_button(
                "⬇️ Download your CV (DOCX)",
                data=st.session_state._last_docx,
                file_name=f"{(st.session_state.profile.get('name') or 'candidate').replace(' ', '_')}_CV.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # HTML preview download
        if download_col2.button("📄 Build & Download HTML preview"):
            name = escape(st.session_state.profile.get("name", ""))
            title = escape(st.session_state.profile.get("title", ""))
            summary = escape(st.session_state.profile.get("summary", "")).replace("\n", "<br>")
            html = f"""<!doctype html><html><head><meta charset="utf-8"><title>{name} — CV</title></head><body>
            <h1>{name}</h1><h3>{title}</h3><p>{summary}</p>"""
            # experience
            if st.session_state.experience_list:
                html += "<h2>Experience</h2>"
                for ex in st.session_state.experience_list:
                    html += f"<h3>{escape(ex.get('role',''))} — {escape(ex.get('company',''))}</h3>"
                    html += "<ul>"
                    for b in ex.get("bullets", []):
                        html += f"<li>{escape(b)}</li>"
                    html += "</ul>"
            if st.session_state.education_list:
                html += "<h2>Education</h2>"
                for ed in st.session_state.education_list:
                    html += f"<div><b>{escape(ed.get('degree',''))}</b> — {escape(ed.get('school',''))} ({escape(ed.get('year',''))})</div>"
            if st.session_state.skills:
                html += "<h2>Skills</h2><div>" + escape(", ".join(st.session_state.skills)) + "</div>"
            html += "</body></html>"
            b = html.encode("utf-8")
            st.download_button("⬇️ Download HTML", data=b, file_name=f"{(st.session_state.profile.get('name') or 'candidate').replace(' ','_')}_CV.html", mime="text/html")

# --------------------------
# Right column: Live preview (HTML styled)
# --------------------------
with right:
    st.subheader("Live preview")

    def render_preview_html(ctx):
        style = st.session_state.design.get("style", "Modern Color")
        accent = st.session_state.design.get("accent", "#0b6efd")
        photo = ctx.get("photo")
        # Modern template
        if style == "Modern Color":
            html = f"""
            <html><head><meta charset='utf-8'>
            <style>
            body{{font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial; color:#1b2b3a; background:transparent;}}
            .card{{background:white;padding:18px;border-radius:10px;box-shadow:0 8px 24px rgba(12,30,60,0.08);width: 380px;}}
            .header{{display:flex;gap:12px;align-items:center;border-bottom:3px solid #f3f6fb;padding-bottom:10px;margin-bottom:10px}}
            .name{{font-weight:700;font-size:20px;color:{accent}}}
            .title{{font-size:13px;color:#334155;margin-top:3px}}
            .meta{{font-size:12px;color:#64748b;margin-top:8px}}
            .section-title{{font-weight:700;margin-top:12px;color:#0f172a;font-size:12px;border-bottom:1px solid #eef2ff;padding-bottom:4px}}
            .skill{{display:inline-block;padding:6px 8px;border-radius:10px;background:#f0faff;margin:4px 4px 0 0;font-size:12px;color:{accent}}}
            ul{{margin:6px 0 0 18px;padding:0}}
            li{{margin-bottom:6px;font-size:13px;}}
            .photo{{width:64px;height:64px;border-radius:8px;object-fit:cover}}
            </style></head><body>
            <div class='card'>
              <div class='header'>
                {f"<img src='{photo}' class='photo'/>" if photo else ""}
                <div>
                  <div class='name'>{ctx.get('name','')}</div>
                  <div class='title'>{ctx.get('title','')}</div>
                  <div class='meta'>{ctx.get('email','')} · {ctx.get('phone','')}</div>
                </div>
              </div>
              <div class='summary'>{ctx.get('summary','')}</div>
              <div class='section-title'>Experience</div>
            """
            for ex in ctx.get("experience", []):
                html += f"<div style='margin-top:8px'><strong>{escape(ex.get('role',''))}</strong> — {escape(ex.get('company',''))} <div style='color:#64748b;font-size:12px'>{escape(ex.get('period',''))}</div>"
                html += "<ul>"
                for b in ex.get("bullets", []):
                    html += f"<li>{escape(b)}</li>"
                html += "</ul></div>"
            html += "<div class='section-title'>Education</div>"
            for ed in ctx.get("education", []):
                html += f"<div style='margin-top:8px'><strong>{escape(ed.get('degree',''))}</strong> — {escape(ed.get('school',''))} <div style='font-size:12px;color:#64748b'>{escape(ed.get('year',''))}</div></div>"
            html += "<div class='section-title'>Skills</div><div>"
            for s in ctx.get("skills", []):
                html += f"<span class='skill'>{escape(s)}</span>"
            html += "</div></div></body></html>"
            return html
        else:
            # Classic / Minimal templates
            html = f"""
            <html><head><meta charset='utf-8'>
            <style>
            body{{font-family: 'Times New Roman', Times, serif; color:#000;}}
            .paper{{width:380px;padding:16px;background:white;border:1px solid #eee}}
            h1{{margin:0;font-size:20px}}
            h2{{margin:4px 0 8px 0;font-size:13px;color:#333}}
            .muted{{color:#555;font-size:12px}}
            ul{{margin:6px 0 0 18px;padding:0}}
            li{{margin-bottom:6px;font-size:13px}}
            </style></head><body><div class='paper'>
            <h1>{escape(ctx.get('name',''))}</h1><div class='muted'>{escape(ctx.get('title',''))}</div>
            <div style='margin-top:8px'>{escape(ctx.get('summary',''))}</div>
            <h2>Experience</h2>
            """
            for ex in ctx.get("experience", []):
                html += f"<div><strong>{escape(ex.get('role',''))}</strong> — {escape(ex.get('company',''))} <div class='muted'>{escape(ex.get('period',''))}</div>"
                html += "<ul>"
                for b in ex.get("bullets", []):
                    html += f"<li>{escape(b)}</li>"
                html += "</ul></div>"
            html += "<h2>Education</h2>"
            for ed in ctx.get("education", []):
                html += f"<div><strong>{escape(ed.get('degree',''))}</strong> — {escape(ed.get('school',''))} ({escape(ed.get('year',''))})</div>"
            html += "<h2>Skills</h2><div>" + escape(", ".join(ctx.get("skills", []))) + "</div>"
            html += "</div></body></html>"
            return html

    # build context from session state
    preview_ctx = {
        "name": st.session_state.profile.get("name", ""),
        "title": st.session_state.profile.get("title", ""),
        "summary": st.session_state.profile.get("summary", ""),
        "email": st.session_state.profile.get("email", ""),
        "phone": st.session_state.profile.get("phone", ""),
        "experience": st.session_state.experience_list,
        "education": st.session_state.education_list,
        "skills": st.session_state.skills,
        "photo": pil_to_datauri(st.session_state.photo_file) if (st.session_state.photo_file and st.session_state.design.get("include_photo", True)) else None
    }

    preview_html = render_preview_html(preview_ctx)
    st.components.v1.html(preview_html, height=720, scrolling=True)

# --------------------------
# Footer / tips
# --------------------------
st.markdown("---")
st.markdown("**Tips:** Edit bullets to add real numbers. Use the sample data button to see how the generator works. The exported `.docx` is fully editable in Word or Google Docs.")
